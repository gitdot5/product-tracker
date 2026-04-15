"""
Stage 5 — MedSum-style Medical Chronology (.docx).

Produces the main deliverable matching MedSum's "Medical Chronology/Summary"
document. Accepts a `ChronologyDoc` JSON (see `stage5_schema.py`) or plain
dict; emits a .docx with all 8 required sections in the exact order observed
across the 8-case corpus:

    1. Title + confidentiality notice
    2. Usage Guidelines/Instructions (verbatim boilerplate)
    3. General Instructions (3-4 patient-specific bullets)
    4. Injury Report (2-col table: DESCRIPTION | DETAILS)
    5. Flow of events (date-range summaries grouped by provider)
    6. Patient History (exactly 5 lines)
    7. Detailed Summary (4-col table: DATE | FACILITY | MEDICAL EVENTS | PDF REF)

MedSum ships `.doc` (Word 97-2003 binary); we produce `.docx` which the
downstream UNC pipeline accepts. If binary `.doc` is required, convert with
`libreoffice --headless --convert-to doc` post-hoc.

Usage:
    python -m pipeline.stage5_chronology_docx \\
        --input  chronology.json \\
        --output "Medical Chronology - Patient.docx"
"""

from __future__ import annotations

import argparse
import json
import logging
from pathlib import Path
from typing import Union

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pipeline.stage5_schema import (
    ChronologyDoc, PatientHistory, Encounter, from_dict,
)

log = logging.getLogger(__name__)


# ── Styling ────────────────────────────────────────────────────────────────

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
TITLE_SIZE = Pt(14)
SECTION_SIZE = Pt(12)

RED = RGBColor(0xC0, 0x00, 0x00)        # reviewer comments
YELLOW_HIGHLIGHT = "FFFF00"              # case-significant details
HEADER_SHADE = "D9D9D9"                  # table header grey


# ── Boilerplate (verbatim from observed MedSum cases) ──────────────────────

USAGE_GUIDELINES = [
    ("Verbatim summary",
     "All the medical details have been included \u201cword by word\u201d or \u201cas it is\u201d "
     "from the provided medical records to avoid alteration of the meaning and to "
     "maintain the validity of the medical records. The sentence available in the "
     "medical record will be taken as it is without any changes to the tense."),
    ("Case synopsis/Flow of events",
     "For ease of reference and to know the glimpse of the case, we have provided "
     "a brief summary including the significant case details."),
    ("Injury report",
     "Injury report outlining the significant medical events/injuries is provided "
     "which will give a general picture of the case."),
    ("Comments",
     "We have included comments for any noteworthy communications, contradictory "
     "information, discrepancies, misinterpretation, missing records, "
     "clarifications, etc for your notification and understanding. The comments "
     "will appear in red italics as follows: \u201c*Reviewer\u2019s Comments\u201d."),
    ("Indecipherable notes/date",
     "Illegible and missing dates are presented as \u201c00/00/0000\u201d (mm/dd/yyyy format). "
     "Illegible handwritten notes are left as a blank space \u201c_____\u201d with a note as "
     "\u201cIllegible Notes\u201d in heading reference."),
    ("Patient\u2019s History",
     "Pre-existing history of the patient has been included in the history section."),
    ("Snapshot inclusion",
     "If the provider name is not decipherable, then the snapshot of the signature "
     "is included. Snapshots of significant examinations and pictorial "
     "representation have been included for reference."),
    ("De-Duplication",
     "Duplicate records and repetitive details have been excluded."),
]

# MedSum sub-headings rendered bold inside Detailed Summary cells
BOLD_SUBHEADINGS = {
    "Chief complaint", "History of present illness", "Presentation",
    "Review of systems", "Vital signs", "Physical exam", "Physical Examination",
    "Differential diagnosis", "MDM notes", "Clinical impression",
    "Assessment", "Plan", "Medications", "Imaging", "Labs", "Radiology",
    "Past medical history", "Surgical history", "Family history",
    "Social history", "Allergy", "Allergies",
    "Current medications", "Current symptoms", "Objective", "Subjective",
    "Home instructions", "Discharge instructions", "Recommendations",
}


# ── Styling helpers ────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _format_run(run, *, bold=False, italic=False, size=BODY_SIZE,
                color: RGBColor = None, highlight: str = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if highlight:
        # w:highlight color name (yellow/green/red etc). Docx only supports
        # named colors for character-level highlight; hex is used via
        # shading on runs which is more fiddly. "yellow" covers our need.
        rPr = run._element.get_or_add_rPr()
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "yellow")
        rPr.append(hl)


def _para(doc, text: str = "", *, bold=False, italic=False,
          size=BODY_SIZE, color=None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _format_run(run, bold=bold, italic=italic, size=size, color=color)


def _cell_text(cell, text: str, *, bold=False, italic=False,
               color=None) -> None:
    cell.text = ""  # clear default empty paragraph
    p = cell.paragraphs[0]
    run = p.add_run(text)
    _format_run(run, bold=bold, italic=italic, color=color)


def _cell_bullets(cell, items, *, italic=False) -> None:
    """Render an unordered list (one • item per line) in a cell."""
    cell.text = ""
    first = True
    for item in items:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        run = p.add_run(f"\u2022 {item}")
        _format_run(run, italic=italic)


# ── Section writers ────────────────────────────────────────────────────────

def _write_title(doc) -> None:
    _para(doc, "Medical Chronology/Summary", bold=True, size=TITLE_SIZE)
    doc.add_paragraph("")
    _para(doc, "Confidential and privileged information", italic=True)
    doc.add_paragraph("")


def _write_usage_guidelines(doc) -> None:
    _para(doc, "Usage guideline/Instructions", bold=True, size=SECTION_SIZE)
    for label, body in USAGE_GUIDELINES:
        p = doc.add_paragraph()
        run = p.add_run(f"*{label}: ")
        _format_run(run, bold=True)
        run2 = p.add_run(body)
        _format_run(run2)
    doc.add_paragraph("")


def _write_general_instructions(doc, d: ChronologyDoc) -> None:
    _para(doc, "General Instructions:", bold=True, size=SECTION_SIZE)
    doc.add_paragraph("")
    # Defaults match MedSum's typical bullets
    bullets = list(d.general_instructions) if d.general_instructions else []
    if not bullets:
        dois = " and ".join(d.injury_report.dates_of_injury) or "[DOI]"
        bullets = [
            f"The medical summary focuses on {d.injury_report.incident_type} on "
            f"{dois}, the injuries and clinical condition of {d.patient.name} as "
            f"a result of the injury, treatments rendered for the complaints and "
            f"progress of the condition.",
            "Initial and final therapy evaluation has been summarized in detail. "
            "Interim visits have been presented cumulatively to avoid repetition "
            "and for ease of reference.",
            "Prior related records have been summarized in detail and prior "
            "unrelated records have been summarized in brief manner.",
        ]
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(b)
        _format_run(run)
    doc.add_paragraph("")


def _write_injury_report(doc, d: ChronologyDoc) -> None:
    _para(doc, "Injury Report:", bold=True, size=SECTION_SIZE)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _cell_text(hdr[0], "DESCRIPTION", bold=True)
    _cell_text(hdr[1], "DETAILS", bold=True)
    for c in hdr:
        _set_cell_shading(c, HEADER_SHADE)

    ir = d.injury_report

    # Row 1: Prior injury details
    row = table.add_row().cells
    _cell_text(row[0], "Prior injury details", bold=True)
    if ir.prior_injury_details:
        _cell_bullets(row[1], ir.prior_injury_details)
    else:
        _cell_text(row[1], "None available")

    # Row 2: Date(s) of injury
    row = table.add_row().cells
    label = "Date of injury" if len(ir.dates_of_injury) <= 1 else "Date of injuries"
    _cell_text(row[0], label, bold=True)
    _cell_text(row[1], " and ".join(ir.dates_of_injury) or "")

    # Row 3: Description
    row = table.add_row().cells
    label = "Description of injury" if len(ir.dates_of_injury) <= 1 else "Description of injuries"
    _cell_text(row[0], label, bold=True)
    _cell_text(row[1], ir.description or "")

    # Row 4: Diagnoses (bullet list)
    row = table.add_row().cells
    _cell_text(row[0], "Injuries/ Diagnoses", bold=True)
    if ir.diagnoses:
        _cell_bullets(row[1], ir.diagnoses)
    else:
        _cell_text(row[1], "")

    # Row 5: Treatments rendered — grouped sub-headings
    row = table.add_row().cells
    _cell_text(row[0], "Treatments rendered", bold=True)
    row[1].text = ""
    groups = [
        ("Medications:", ir.treatments.medications),
        ("Procedures:", ir.treatments.procedures),
        ("Therapy:", ir.treatments.therapy),
        ("Imaging:", ir.treatments.imaging),
        ("Labs:", ir.treatments.labs),
    ]
    first = True
    for heading, items in groups:
        if not items:
            continue
        if first:
            p = row[1].paragraphs[0]
            first = False
        else:
            p = row[1].add_paragraph()
        run = p.add_run(heading)
        _format_run(run, bold=True)
        for item in items:
            p = row[1].add_paragraph()
            run = p.add_run(f"\u2022 {item}")
            _format_run(run)

    doc.add_paragraph("")


def _write_flow_of_events(doc, d: ChronologyDoc) -> None:
    _para(doc, "Flow of events", bold=True, size=SECTION_SIZE)
    doc.add_paragraph("")

    current_group = None
    for entry in d.flow_of_events:
        if entry.provider_group and entry.provider_group != current_group:
            _para(doc, entry.provider_group, bold=True)
            current_group = entry.provider_group
        # Summary line format: "DATE_RANGE: summary"
        prefix = f"{entry.date_range}: " if entry.date_range else ""
        _para(doc, prefix + entry.summary)
        if entry.reviewer_comment:
            _para(doc, f"*Reviewer\u2019s comment: {entry.reviewer_comment}",
                  italic=True, color=RED)
    doc.add_paragraph("")


def _write_patient_history(doc, d: ChronologyDoc) -> None:
    _para(doc, "Patient History", bold=True, size=SECTION_SIZE)
    doc.add_paragraph("")

    def _line(label: str, hl):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        _format_run(run, bold=True)
        run2 = p.add_run(hl.text if hl.text else "Not available.")
        _format_run(run2)
        if hl.pdf_ref:
            run3 = p.add_run(f" (PDF ref: {hl.pdf_ref})")
            _format_run(run3)

    ph = d.patient_history
    _line("Past Medical History", ph.past_medical)
    _line("Surgical History", ph.surgical)
    _line("Family History", ph.family)
    _line("Social History", ph.social)
    _line("Allergy", ph.allergy)

    doc.add_paragraph("")


def _write_detailed_summary(doc, d: ChronologyDoc) -> None:
    _para(doc, "Detailed Summary", bold=True, size=SECTION_SIZE)
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["DATE", "FACILITY/ PROVIDER", "MEDICAL EVENTS", "PDF REF"]):
        _cell_text(hdr[i], h, bold=True)
        _set_cell_shading(hdr[i], HEADER_SHADE)

    # Widen MEDICAL EVENTS column (rough sizing)
    try:
        widths = [Cm(2.5), Cm(4.0), Cm(11.0), Cm(2.5)]
        for i, w in enumerate(widths):
            for row in table.columns[i].cells:
                row.width = w
    except Exception:
        pass  # widths are advisory

    for enc in d.encounters:
        # Optional group header row spanning all 4 columns
        if enc.group_header:
            row = table.add_row().cells
            # Merge 4 cells via python-docx merge
            merged = row[0].merge(row[1]).merge(row[2]).merge(row[3])
            header_text = enc.group_header_text or (
                f"{enc.facility} / {enc.date}" if enc.facility and enc.date
                else enc.facility or enc.date
            )
            _cell_text(merged, header_text, bold=True, italic=True)
            _set_cell_shading(merged, "F2F2F2")

        row = table.add_row().cells
        _cell_text(row[0], enc.date)
        facility_block = [enc.facility] + list(enc.providers)
        row[1].text = ""
        first = True
        for line in facility_block:
            if not line:
                continue
            if first:
                p = row[1].paragraphs[0]
                first = False
            else:
                p = row[1].add_paragraph()
            run = p.add_run(line)
            _format_run(run)

        # MEDICAL EVENTS cell: render with bold sub-headings
        row[2].text = ""
        _render_medical_events(row[2], enc.medical_events)

        _cell_text(row[3], enc.pdf_ref)


def _render_medical_events(cell, text: str) -> None:
    """Write medical_events text into a table cell, bolding known sub-headings.

    A sub-heading is detected when a paragraph starts with 'XYZ:' where XYZ
    (case-insensitive) is in BOLD_SUBHEADINGS. The heading and colon are
    bold; the rest of the paragraph is body text.
    """
    paragraphs = [p for p in text.split("\n") if p is not None]
    first = True
    for line in paragraphs:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        # Detect "Heading: body"
        if ":" in line:
            head, body = line.split(":", 1)
            if head.strip() in BOLD_SUBHEADINGS or head.strip().title() in BOLD_SUBHEADINGS:
                run = p.add_run(f"{head.strip()}: ")
                _format_run(run, bold=True)
                if body.strip():
                    run2 = p.add_run(body.strip())
                    _format_run(run2)
                continue
        # Reviewer comment detection
        if line.strip().startswith("*Reviewer"):
            run = p.add_run(line)
            _format_run(run, italic=True, color=RED)
            continue
        # Plain line
        run = p.add_run(line)
        _format_run(run)


# ── Public API ─────────────────────────────────────────────────────────────

def build_chronology(spec: Union[ChronologyDoc, dict],
                     output_path: Union[str, Path]) -> Path:
    if isinstance(spec, dict):
        spec = from_dict(spec)
    output_path = Path(output_path)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    _write_title(doc)
    _write_usage_guidelines(doc)
    _write_general_instructions(doc, spec)
    _write_injury_report(doc, spec)
    _write_flow_of_events(doc, spec)
    _write_patient_history(doc, spec)
    _write_detailed_summary(doc, spec)

    doc.save(output_path)
    log.info("Wrote Medical Chronology -> %s", output_path)
    return output_path


# ── CLI ────────────────────────────────────────────────────────────────────

def main() -> int:
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s %(levelname)s %(message)s")
    p = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    p.add_argument("--input", required=True,
                   help="Path to ChronologyDoc JSON")
    p.add_argument("--output", required=True, help="Output .docx path")
    args = p.parse_args()

    with open(args.input) as f:
        spec = json.load(f)
    build_chronology(spec, args.output)
    print(f"Output: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
