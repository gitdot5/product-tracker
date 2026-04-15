"""
Stage 5 — MedSum-style Delivery Note (.docx).

Produces the cover letter that ships alongside the Medical Chronology, Merged
Records, and Hyperlinked Records. Accepts a `ChronologyDoc` JSON (see
`stage5_schema.py`) or plain dict; emits a .docx with MedSum's exact 13-block
structure confirmed across 8 observed cases.

Usage:
    python -m pipeline.stage5_delivery_note \\
        --input  chronology.json \\
        --output "Delivery Note - Patient.docx"
"""

from __future__ import annotations

import argparse
import json
import logging
from pathlib import Path
from typing import Union

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pipeline.stage5_schema import ChronologyDoc, from_dict

log = logging.getLogger(__name__)


# ── Styling helpers ────────────────────────────────────────────────────────

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
TITLE_SIZE = Pt(14)
SECTION_SIZE = Pt(12)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Apply a background color to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _para(doc, text: str = "", *, bold: bool = False, size: Pt = BODY_SIZE,
          italic: bool = False, color: RGBColor = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ── Builders ───────────────────────────────────────────────────────────────

def _write_header(doc, d: ChronologyDoc) -> None:
    _para(doc, f"Delivery Note - {d.patient.name}", bold=True, size=TITLE_SIZE)
    doc.add_paragraph("")  # spacer


def _write_greeting(doc, d: ChronologyDoc) -> None:
    _para(doc, f"Dear {d.patient.contact_first_name},")
    doc.add_paragraph("")


def _write_intro(doc, d: ChronologyDoc) -> None:
    _para(doc,
          f"We have completed the medical records review for {d.patient.name} "
          f"and prepared the medical chronology.")
    doc.add_paragraph("")
    _para(doc,
          "As a free service, we have prepared the hyperlinked medical records "
          "for ease of navigation to refer the source document.")
    doc.add_paragraph("")


def _write_chronology_block(doc, d: ChronologyDoc) -> None:
    _para(doc, "Medical chronology:", bold=True)
    doc.add_paragraph("")

    dates = d.injury_report.dates_of_injury
    if len(dates) == 0:
        pass  # intentionally omit line
    elif len(dates) == 1:
        _para(doc, f"Date of injury: {dates[0]}", bold=True)
    else:
        joined = " and ".join(dates)
        _para(doc, f"Dates of injuries: {joined}", bold=True)
    doc.add_paragraph("")

    _para(doc, "Case focus details:", bold=True)
    doc.add_paragraph("")
    if d.case_focus:
        # Accept pre-formatted paragraphs (blank line = paragraph break)
        for block in d.case_focus.split("\n\n"):
            _para(doc, block.strip())
        doc.add_paragraph("")


def _write_causation(doc, d: ChronologyDoc) -> None:
    if not d.causation_statements:
        return
    _para(doc, "Causation:", bold=True)
    doc.add_paragraph("")
    for s in d.causation_statements:
        _para(doc, f"{s.date}: {s.text}" if s.date else s.text)
    doc.add_paragraph("")


def _write_disability(doc, d: ChronologyDoc) -> None:
    if not d.disability_statements:
        return
    _para(doc, "Disability:", bold=True)
    for s in d.disability_statements:
        _para(doc, f"{s.date}: {s.text}" if s.date else s.text)
    doc.add_paragraph("")


def _write_missing_records(doc, d: ChronologyDoc) -> None:
    _para(doc, "Missing medical records:", bold=True)
    doc.add_paragraph("")

    if d.no_missing_records or not d.missing_records:
        _para(doc, "There are no critical missing medical records.")
        doc.add_paragraph("")
        return

    # 6-column table
    headers = [
        "Date/\nPeriod",
        "Provider/\nFacility",
        "What records are needed",
        "Confirmatory/\nProbable",
        "Statement regarding missing records",
        "PDF Reference",
    ]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_shading(hdr[i], "D9D9D9")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
                run.bold = True
    for mr in d.missing_records:
        row = table.add_row().cells
        row[0].text = mr.date_period
        row[1].text = mr.provider
        row[2].text = mr.records_needed
        row[3].text = mr.confirmatory_or_probable
        row[4].text = mr.statement
        row[5].text = mr.pdf_reference
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = BODY_SIZE
    doc.add_paragraph("")
    _para(doc,
          "Please retrieve us the missing medical records, upon which we will "
          "revise the medical chronology.")
    doc.add_paragraph("")


def _write_merged_records_paragraph(doc) -> None:
    _para(doc, "Merged Medical Records:", bold=True)
    doc.add_paragraph("")
    _para(doc,
          "For ease of reference, we have merged all the medical records together "
          "in the order of receipt and have captured the page number as reference "
          "in the chronology. Kindly refer to the page numbers of the merged "
          "medical records at the left lower corner when referring to details in "
          "the chronology.")
    doc.add_paragraph("")


def _write_hyperlinked_paragraph(doc) -> None:
    _para(doc, "Hyperlinked Medical Records:", bold=True)
    doc.add_paragraph("")
    _para(doc,
          "Place the hand symbol on the page reference and click the page number "
          "to refer the corresponding source document.")
    doc.add_paragraph("")
    _para(doc, "Use Adobe Reader/Adobe Acrobat to navigate the page view.")
    doc.add_paragraph("")
    _para(doc,
          "Keyboard shortcuts for going back and forward: Alt+ Left Arrow or "
          "Alt+ Right Arrow, respectively")
    doc.add_paragraph("")


def _write_closing(doc) -> None:
    _para(doc, "We will be happy to make any modifications if needed.")
    doc.add_paragraph("")
    _para(doc,
          "Please feel free to reach us if you need any further assistance with "
          "the files and kindly acknowledge the receipt of the files. Thanks!")
    doc.add_paragraph("")
    _para(doc, "*****")


# ── Public API ─────────────────────────────────────────────────────────────

def build_delivery_note(spec: Union[ChronologyDoc, dict],
                        output_path: Union[str, Path]) -> Path:
    if isinstance(spec, dict):
        spec = from_dict(spec)
    output_path = Path(output_path)

    doc = Document()
    # Set default font for the document
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE

    _write_header(doc, spec)
    _write_greeting(doc, spec)
    _write_intro(doc, spec)
    _write_chronology_block(doc, spec)
    _write_causation(doc, spec)
    _write_disability(doc, spec)
    _write_missing_records(doc, spec)
    _write_merged_records_paragraph(doc)
    _write_hyperlinked_paragraph(doc)
    _write_closing(doc)

    doc.save(output_path)
    log.info("Wrote Delivery Note -> %s", output_path)
    return output_path


# ── CLI ────────────────────────────────────────────────────────────────────

def main() -> int:
    logging.basicConfig(level=logging.INFO,
                        format="%(asctime)s %(levelname)s %(message)s")
    p = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    p.add_argument("--input", required=True,
                   help="Path to ChronologyDoc JSON")
    p.add_argument("--output", required=True, help="Output .docx path")
    args = p.parse_args()

    with open(args.input) as f:
        spec = json.load(f)
    build_delivery_note(spec, args.output)
    print(f"Output: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
